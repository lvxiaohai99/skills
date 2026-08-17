#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""从简历文件提取纯文本，供面试小助手后续分析。

支持: pdf / docx / doc / xlsx / xls / md / txt / html / csv
用法: python extract_resume.py <文件路径>
"""

from __future__ import annotations

import csv
import json
import os
import sys
import traceback
from html.parser import HTMLParser
from typing import List, Tuple


def _force_utf8_stdio() -> None:
    """Windows 终端默认代码页常把中文打成乱码；强制 stdout/stderr 用 UTF-8。"""
    os.environ.setdefault("PYTHONIOENCODING", "utf-8")
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass


_force_utf8_stdio()

SUPPORTED = {
    ".pdf",
    ".docx",
    ".doc",
    ".xlsx",
    ".xls",
    ".md",
    ".markdown",
    ".txt",
    ".html",
    ".htm",
    ".csv",
}


def fail(message: str, code: int = 1) -> None:
    """把错误打到 stderr 并退出，方便在对话里定位失败步骤。"""
    print(f"[extract_resume] ERROR: {message}", file=sys.stderr)
    sys.exit(code)


def read_text_file(path: str) -> str:
    """按常见中文编码依次尝试读取纯文本。"""
    raw = open(path, "rb").read()
    for enc in ("utf-8-sig", "utf-8", "gb18030", "gbk", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def extract_pdf(path: str) -> Tuple[str, List[str]]:
    warnings: List[str] = []
    try:
        import pdfplumber
    except ImportError:
        fail("缺少 pdfplumber。请执行: pip install pdfplumber")
    chunks: List[str] = []
    with pdfplumber.open(path) as pdf:
        if not pdf.pages:
            warnings.append("empty_pdf")
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            tables = page.extract_tables() or []
            table_lines: List[str] = []
            for table in tables:
                for row in table:
                    cells = [str(c).strip() if c is not None else "" for c in row]
                    if any(cells):
                        table_lines.append(" | ".join(cells))
            page_body = text.strip()
            if table_lines:
                page_body = (page_body + "\n" + "\n".join(table_lines)).strip()
            if page_body:
                chunks.append(f"--- 第 {i} 页 ---\n{page_body}")
    if not chunks:
        warnings.append("low_text")
        warnings.append("可能是扫描件 PDF，无法提取文字")
    return "\n\n".join(chunks), warnings


def extract_docx(path: str) -> Tuple[str, List[str]]:
    try:
        from docx import Document
    except ImportError:
        fail("缺少 python-docx。请执行: pip install python-docx")
    doc = Document(path)
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    warnings = ["low_text"] if len(text.strip()) < 200 else []
    return text, warnings


def extract_doc(path: str) -> Tuple[str, List[str]]:
    """旧版 .doc：优先用 Word COM（Windows），失败则提示转成 docx。"""
    warnings: List[str] = ["doc_legacy"]
    try:
        import win32com.client  # type: ignore
    except ImportError:
        fail(
            "当前是旧版 .doc。请另存为 .docx 后重试；"
            "或安装 pywin32 后在已安装 Word 的 Windows 上再跑本脚本。"
        )
    word = None
    try:
        abs_path = os.path.abspath(path)
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(abs_path, ReadOnly=True)
        text = doc.Content.Text or ""
        doc.Close(False)
        warnings.append("extracted_via_word_com")
        if len(text.strip()) < 200:
            warnings.append("low_text")
        return text.replace("\r", "\n"), warnings
    except Exception as exc:
        fail(f"无法读取 .doc（{exc}）。请用 Word 另存为 .docx 后重试。")
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
    return "", warnings


def _sheet_to_text(rows: List[List[object]], sheet_name: str) -> str:
    lines = [f"--- 工作表: {sheet_name} ---"]
    for row in rows:
        cells = ["" if c is None else str(c).strip() for c in row]
        if any(cells):
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_xlsx(path: str) -> Tuple[str, List[str]]:
    try:
        import openpyxl
    except ImportError:
        fail("缺少 openpyxl。请执行: pip install openpyxl")
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    chunks = []
    for name in wb.sheetnames:
        ws = wb[name]
        rows = [list(row) for row in ws.iter_rows(values_only=True)]
        chunks.append(_sheet_to_text(rows, name))
    wb.close()
    text = "\n\n".join(chunks)
    warnings = ["low_text"] if len(text.strip()) < 200 else []
    return text, warnings


def extract_xls(path: str) -> Tuple[str, List[str]]:
    try:
        import xlrd
    except ImportError:
        fail("缺少 xlrd。请执行: pip install xlrd")
    book = xlrd.open_workbook(path)
    chunks = []
    for sheet in book.sheets():
        rows = [sheet.row_values(r) for r in range(sheet.nrows)]
        chunks.append(_sheet_to_text(rows, sheet.name))
    text = "\n\n".join(chunks)
    warnings = ["low_text"] if len(text.strip()) < 200 else []
    return text, warnings


class _HTMLText(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._skip = False
        self.parts: List[str] = []

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in ("script", "style"):
            self._skip = True

    def handle_endtag(self, tag: str) -> None:
        if tag in ("script", "style"):
            self._skip = False
        if tag in ("p", "div", "br", "li", "tr", "h1", "h2", "h3", "h4"):
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self._skip and data.strip():
            self.parts.append(data.strip())


def extract_html(path: str) -> Tuple[str, List[str]]:
    parser = _HTMLText()
    parser.feed(read_text_file(path))
    text = " ".join(parser.parts)
    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    warnings = ["low_text"] if len(text) < 200 else []
    return text, warnings


def extract_csv(path: str) -> Tuple[str, List[str]]:
    raw = read_text_file(path)
    rows = list(csv.reader(raw.splitlines()))
    lines = [" | ".join(c.strip() for c in row) for row in rows if any(x.strip() for x in row)]
    text = "\n".join(lines)
    warnings = ["low_text"] if len(text) < 200 else []
    return text, warnings


def extract(path: str) -> Tuple[str, str, List[str]]:
    ext = os.path.splitext(path)[1].lower()
    if ext not in SUPPORTED:
        fail(f"不支持的格式: {ext}。支持: {', '.join(sorted(SUPPORTED))}")
    if ext == ".pdf":
        text, warnings = extract_pdf(path)
    elif ext == ".docx":
        text, warnings = extract_docx(path)
    elif ext == ".doc":
        text, warnings = extract_doc(path)
    elif ext == ".xlsx":
        text, warnings = extract_xlsx(path)
    elif ext == ".xls":
        text, warnings = extract_xls(path)
    elif ext in {".html", ".htm"}:
        text, warnings = extract_html(path)
    elif ext == ".csv":
        text, warnings = extract_csv(path)
    else:
        text, warnings = read_text_file(path), []
        if len(text.strip()) < 200:
            warnings.append("low_text")
    return ext.lstrip("."), text.strip(), warnings


def main() -> None:
    if len(sys.argv) < 2:
        fail("用法: python extract_resume.py <简历文件路径>")
    path = sys.argv[1]
    if not os.path.isfile(path):
        fail(f"文件不存在: {path}")
    try:
        fmt, text, warnings = extract(path)
    except SystemExit:
        raise
    except Exception:
        traceback.print_exc()
        fail("解析失败，见上方堆栈。可改用 .docx / .md / .txt 后重试。")
    meta = {
        "file": os.path.abspath(path),
        "format": fmt,
        "chars": len(text),
        "warnings": warnings,
    }
    print("===RESUME_META===")
    print(json.dumps(meta, ensure_ascii=False, indent=2))
    print("===RESUME_TEXT===")
    print(text if text else "(无文本)")


if __name__ == "__main__":
    main()
