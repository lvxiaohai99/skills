#!/usr/bin/env python3
"""Clone a Word 详细设计模板并按同一套样式重建正文。

Agent 写一个薄封装脚本 import 本模块，不要从零用 docx-js 重做封面/页眉页脚。

典型用法：

    from sdd_docx import ReportBuilder
    b = ReportBuilder(template, output)
    b.open()
    b.cover(title="G200Z软件详细设计报告_OTA模块", version="V1.0.0")
    b.revision_page([["2026-08-18", "V1.0.0", "", "全部", "根据代码首次编制"]])
    b.toc_page("目 录", [("1  概要", 4, 1), ("1.1  目的", 4, 2)])
    b.heading("概要", 1)
    b.body("……")
    b.save()
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

SONG = "宋体"
TIMES = "Times New Roman"
TBL_W = 10774  # 与中信科 V1.0.2 函数表宽度对齐


def insert_before(parent, child, *tag_names):
    for name in tag_names:
        sib = parent.find(qn(name))
        if sib is not None:
            sib.addprevious(child)
            return
    parent.append(child)


class ReportBuilder:
    def __init__(self, template: str | Path, output: str | Path):
        self.template = Path(template)
        self.output = Path(output)
        self.doc = None

    def open(self) -> None:
        shutil.copy(self.template, self.output)
        self.doc = Document(str(self.output))
        body = self.doc.element.body
        sectPr = body.find(qn("w:sectPr"))
        for child in list(body):
            if child is not sectPr:
                body.remove(child)

    def save(self) -> None:
        self.doc.core_properties.title = self.output.stem
        self.doc.save(str(self.output))

    # ----- runs / paragraphs -----

    def set_run_font(self, run, east=SONG, ascii_font=TIMES, size=None, bold=None):
        run.font.name = ascii_font
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), ascii_font)
        rFonts.set(qn("w:hAnsi"), ascii_font)
        rFonts.set(qn("w:eastAsia"), east)
        rFonts.set(qn("w:cs"), ascii_font)
        if size is not None:
            run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold

    def add_text(self, p, text, *, east=SONG, ascii_font=TIMES, size=None, bold=None):
        run = p.add_run(text)
        self.set_run_font(run, east=east, ascii_font=ascii_font, size=size, bold=bold)
        return run

    def empty(self, n=1, align=None):
        for _ in range(n):
            p = self.doc.add_paragraph()
            p.style = self.doc.styles["Normal"]
            if align is not None:
                p.alignment = align
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

    def page_break(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)

    def heading(self, text, level: int):
        style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}[level]
        p = self.doc.add_paragraph(style=style)
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        size = {1: 15, 2: 14, 3: 12, 4: 12}[level]
        ascii_font = TIMES if level == 1 else SONG
        self.add_text(p, text, east=SONG, ascii_font=ascii_font, size=size, bold=True)
        return p

    def body(self, text, *, justify=True):
        p = self.doc.add_paragraph()
        p.style = self.doc.styles["Normal"]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        self.add_text(p, text, size=12)
        return p

    def note(self, text):
        p = self.doc.add_paragraph()
        p.style = self.doc.styles["Normal"]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        self.add_text(p, text, size=12, bold=True)
        return p

    def caption(self, text):
        p = self.doc.add_paragraph(style="Caption")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        self.add_text(p, text, east="黑体", ascii_font=TIMES, size=10.5)
        return p

    def figure(self, path, cap, width_cm=15.4):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(str(path), width=Cm(width_cm))
        self.caption(cap)

    # ----- tables (OOXML child order matters for XSD) -----

    def _shade(self, cell, fill="FFFFFF"):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            insert_before(tcPr, shd, "w:noWrap", "w:tcMar", "w:textDirection",
                          "w:tcFitText", "w:vAlign", "w:hideMark")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)

    def _cell_width(self, cell, dxa):
        tcPr = cell._tc.get_or_add_tcPr()
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None:
            tcW = OxmlElement("w:tcW")
            tcPr.append(tcW)
        tcW.set(qn("w:w"), str(dxa))
        tcW.set(qn("w:type"), "dxa")

    def _fill_cell(self, cell, text, *, bold=False, size=10.5, align="center"):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.clear()
        p.alignment = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        self.add_text(p, text, size=size, bold=bold)

    def _set_grid(self, tbl, widths):
        tblPr = tbl._tbl.tblPr
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), str(sum(widths)))
        tblW.set(qn("w:type"), "dxa")
        grid = tbl._tbl.find(qn("w:tblGrid"))
        if grid is None:
            grid = OxmlElement("w:tblGrid")
            tbl._tbl.insert(1, grid)
        for child in list(grid):
            grid.remove(child)
        for w in widths:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(w))
            grid.append(gc)
        for row in tbl.rows:
            for cell, w in zip(row.cells, widths):
                self._cell_width(cell, w)

    def _grid_borders(self, tbl):
        tblPr = tbl._tbl.tblPr
        borders = tblPr.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            insert_before(tblPr, borders, "w:shd", "w:tblLayout", "w:tblCellMar",
                          "w:tblLook", "w:tblCaption", "w:tblDescription")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = borders.find(qn(f"w:{edge}"))
            if el is None:
                el = OxmlElement(f"w:{edge}")
                borders.append(el)
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")

    def grid_table(self, headers, rows, widths, header_fill="F2F2F2"):
        tbl = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._grid_borders(tbl)
        self._set_grid(tbl, widths)
        for i, h in enumerate(headers):
            self._fill_cell(tbl.rows[0].cells[i], h, bold=True)
            self._shade(tbl.rows[0].cells[i], header_fill)
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                self._fill_cell(tbl.rows[r + 1].cells[c], val,
                                align="center" if c == 0 else "left")
        return tbl

    def kv_table(self, pairs, widths=(2400, 8374)):
        """函数说明表：左列标签、右列内容。"""
        tbl = self.doc.add_table(rows=len(pairs), cols=2)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._grid_borders(tbl)
        self._set_grid(tbl, list(widths))
        for i, (k, v) in enumerate(pairs):
            self._fill_cell(tbl.rows[i].cells[0], k, bold=True, align="center")
            self._shade(tbl.rows[i].cells[0], "F2F2F2")
            self._fill_cell(tbl.rows[i].cells[1], v, align="left")
        return tbl

    def func_spec(self, name, ret, params, purpose, calls, decl, impl):
        return self.kv_table([
            ("函数名", name),
            ("函数返回值", ret),
            ("参数列表", params),
            ("函数功能", purpose),
            ("调用的函数", calls),
            ("函数定义文件", decl),
            ("函数实现文件", impl),
        ])

    # ----- cover / revision / TOC (中信科 V1.0.2 版式) -----

    def cover(self, title: str, version: str, company="中信科智联科技有限公司",
              release_date="", authors=("", "", ""),
              legal=None):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_text(p, title, east=SONG, ascii_font=SONG, size=28, bold=True)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_text(p, f"（{version}）", east=SONG, ascii_font=SONG, size=22)

        self.empty(2, WD_ALIGN_PARAGRAPH.CENTER)
        self._author_table(authors, release_date)
        self.empty(6, WD_ALIGN_PARAGRAPH.JUSTIFY)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.add_text(p, "版权所有", east=SONG, ascii_font=SONG, size=12)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.add_text(p, company, east=SONG, ascii_font=SONG, size=12, bold=True)

        if legal is None:
            legal = (
                f"本资料及其包含的所有内容为{company}（简称中信科智联）所有,受中国法律及适用之国际公约中有关著作权法律的保护。"
                "未经中信科智联书面授权，任何人不得以任何形式复制、传播、散布、改动或以其它方式使用本资料的部分或全部内容，违者将被依法追究责任。"
            )
        p = self.doc.add_paragraph()
        self.add_text(p, legal, east=SONG, ascii_font=SONG, size=11)
        run = p.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)

    def _author_table(self, authors, release_date):
        tbl = self.doc.add_table(rows=4, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_grid(tbl, (1418, 2867))
        tblPr = tbl._tbl.tblPr
        jc = tblPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            tblPr.append(jc)
        jc.set(qn("w:val"), "center")
        labels = ["编 写 者", "审 核 者", "审 批 者", "发布日期"]
        values = [authors[0] if len(authors) > 0 else "",
                  authors[1] if len(authors) > 1 else "",
                  authors[2] if len(authors) > 2 else "",
                  release_date]
        for i, (lab, val) in enumerate(zip(labels, values)):
            row = tbl.rows[i]
            trPr = row._tr.get_or_add_trPr()
            h = OxmlElement("w:trHeight")
            h.set(qn("w:val"), "488")
            trPr.append(h)
            left, right = row.cells
            for c in (left, right):
                c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                self._shade(c, "FFFFFF")
            self._nil_borders(left)
            self._bottom_only(right)
            self._fill_cell(left, lab, size=12, align="both")
            self._fill_cell(right, val, size=12, align="center")

    def _nil_borders(self, cell):
        self._set_borders(cell, top="nil", left="nil", bottom="nil", right="nil")

    def _bottom_only(self, cell):
        self._set_borders(cell, top="nil", left="nil",
                          bottom=("single", "4", "auto"), right="nil")

    def _set_borders(self, cell, **sides):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            insert_before(tcPr, tcBorders, "w:shd", "w:noWrap", "w:tcMar",
                          "w:textDirection", "w:tcFitText", "w:vAlign", "w:hideMark")
        for edge, spec in sides.items():
            el = tcBorders.find(qn(f"w:{edge}"))
            if el is None:
                el = OxmlElement(f"w:{edge}")
                tcBorders.append(el)
            if spec == "nil":
                el.set(qn("w:val"), "nil")
            else:
                val, sz, color = spec
                el.set(qn("w:val"), val)
                el.set(qn("w:sz"), str(sz))
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), color)

    def revision_page(self, rows):
        p = self.doc.add_paragraph()
        self.add_text(p, "修订记录", east=SONG, ascii_font=SONG, size=16)
        self.grid_table(
            ["日期", "版本", "拟制/修订人", "修订章节", "修订内容"],
            rows,
            [1800, 1400, 2000, 1600, 3974],
        )
        self.empty(1)
        self.page_break()

    def toc_page(self, title, entries):
        """entries: list of (text, page, level) level=1..4。行距压紧以免溢出成空白页。"""
        p = self.doc.add_paragraph(style="TOC")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        self.add_text(p, title, east=SONG, ascii_font=SONG, size=16, bold=True)
        for text, page, level in entries:
            style = {1: "toc 1", 2: "toc 2", 3: "toc 3", 4: "toc 4"}[level]
            para = self.doc.add_paragraph(style=style)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            self.add_text(para, text, size=12, bold=(level == 1))
            run = para.add_run()
            ptab = OxmlElement("w:ptab")
            ptab.set(qn("w:alignment"), "right")
            ptab.set(qn("w:relativeTo"), "margin")
            ptab.set(qn("w:leader"), "dot")
            run._r.append(ptab)
            self.add_text(para, str(page), size=12)
        self.page_break()
