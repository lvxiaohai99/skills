#!/usr/bin/env python3
"""Dump a Word 详细设计模板的结构，供后续按同一套样式填内容。

Usage:
    python inspect_template.py <template.docx> [-o structure.json]
"""
from __future__ import annotations

import argparse
import json
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _text(el) -> str:
    return "".join(el.xpath(".//w:t/text()", namespaces=NS))


def inspect(path: Path) -> dict:
    doc = Document(str(path))
    sec = doc.sections[0]
    styles_used = Counter()
    headings = []
    writer_notes = []
    paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        name = p.style.name if p.style else ""
        text = p.text.strip()
        styles_used[name] += 1
        if name.startswith("Heading") and text:
            headings.append({"index": i, "style": name, "text": text})
        if "【模板说明】" in text or text.startswith("【示例"):
            writer_notes.append({"index": i, "text": text})
        if text:
            paragraphs.append({"index": i, "style": name, "text": text[:200]})

    tables = []
    for i, t in enumerate(doc.tables):
        first = [c.text.replace("\n", " ").strip()[:40] for c in t.rows[0].cells] if t.rows else []
        tables.append({
            "index": i,
            "rows": len(t.rows),
            "cols": len(t.columns) if t.rows else 0,
            "header": first,
        })

    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml")
        has_toc = b"TOC " in xml or b"TOC\\o" in xml
        root = etree.fromstring(xml)

    body_order = []
    tbl_i = 0
    p_i = 0
    body = root.find("w:body", NS)
    for child in body:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            t = _text(child)[:60]
            brs = [b.get(W + "type") for b in child.xpath(".//w:br", namespaces=NS)]
            body_order.append({"kind": "p", "i": p_i, "text": t, "br": brs})
            p_i += 1
        elif tag == "tbl":
            body_order.append({"kind": "tbl", "i": tbl_i, "text": _text(child)[:40]})
            tbl_i += 1

    return {
        "file": str(path),
        "page": {
            "width_cm": round(sec.page_width.cm, 2),
            "height_cm": round(sec.page_height.cm, 2),
            "margin_cm": {
                "left": round(sec.left_margin.cm, 2),
                "right": round(sec.right_margin.cm, 2),
                "top": round(sec.top_margin.cm, 2),
                "bottom": round(sec.bottom_margin.cm, 2),
            },
            "different_first_page": bool(sec.different_first_page_header_footer),
        },
        "has_toc_field": has_toc,
        "styles_used": dict(styles_used.most_common()),
        "headings": headings,
        "writer_notes": writer_notes,
        "tables": tables,
        "inline_shapes": len(doc.inline_shapes),
        "body_order_prefix": body_order[:80],
        "paragraphs_nonempty_prefix": paragraphs[:80],
    }


def main() -> None:
    ap = argparse.ArgumentParser(description="Inspect a detailed-design Word template")
    ap.add_argument("template")
    ap.add_argument("-o", "--output", help="Write JSON to this path (default stdout)")
    args = ap.parse_args()
    data = inspect(Path(args.template))
    text = json.dumps(data, ensure_ascii=False, indent=2)
    if args.output:
        Path(args.output).write_text(text, encoding="utf-8")
        print(f"wrote {args.output}")
    else:
        print(text)


if __name__ == "__main__":
    main()
